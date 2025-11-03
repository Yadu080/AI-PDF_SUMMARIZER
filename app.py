import os
import io
import time
from datetime import datetime
from flask import Flask, render_template, request, jsonify, send_file, session
from flask_socketio import SocketIO, emit
from flask_sqlalchemy import SQLAlchemy
import google.generativeai as genai
import pdfplumber
from PyPDF2 import PdfReader
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt, RGBColor
from google.generativeai import GenerativeModel

# Load environment variables
load_dotenv()

# Initialize Flask app
app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('FLASK_SECRET_KEY', os.urandom(32))
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///instance/history.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['MAX_CONTENT_LENGTH'] = int(os.getenv('MAX_FILE_SIZE_MB', 10)) * 1024 * 1024
genai_api_key = os.environ.get("GEMINI_API_KEY")
if not genai_api_key:
    raise ValueError("Please set GEMINI_API_KEY in your .env file or environment")

genai.configure(api_key=genai_api_key)

# Model ID from env or default
GEMINI_MODEL_ID = os.environ.get("GEMINI_MODEL_ID", "gemini-2.5-flash")
print("Current Working Directory:", os.getcwd())
basedir = os.path.abspath(os.path.dirname(__file__))
instance_path = os.path.join(basedir, 'instance')
os.makedirs(instance_path, exist_ok=True)  # Create folder if missing

db_path = os.path.join(instance_path, 'history.db')
app.config['SQLALCHEMY_DATABASE_URI'] = f"sqlite:///{db_path}"
# Initialize extensions
db = SQLAlchemy(app)
socketio = SocketIO(app, cors_allowed_origins="*", async_mode='eventlet')


# Constants
ALLOWED_EXTENSIONS = {'pdf'}
MAX_TEXT_LENGTH = int(os.getenv('MAX_TEXT_LENGTH', 8000))
UPLOAD_FOLDER = os.path.join('static', 'uploads')

# Ensure upload folder exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs('instance', exist_ok=True)

# ==================== DATABASE MODELS ====================

class SummaryHistory(db.Model):
    """Model to store summary history"""
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255), nullable=False)
    summary_text = db.Column(db.Text, nullable=False)
    summary_type = db.Column(db.String(50), default='standard')
    page_count = db.Column(db.Integer)
    word_count = db.Column(db.Integer)
    rating = db.Column(db.Integer, default=0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    def to_dict(self):
        return {
            'id': self.id,
            'filename': self.filename,
            'summary_text': self.summary_text,
            'summary_type': self.summary_type,
            'page_count': self.page_count,
            'word_count': self.word_count,
            'rating': self.rating,
            'created_at': self.created_at.strftime('%Y-%m-%d %H:%M:%S')
        }

# Create database tables
with app.app_context():
    db.create_all()

# ==================== HELPER FUNCTIONS ====================

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    """
    Extract text from PDF using pdfplumber (primary) and PyPDF2 (fallback)
    Returns: (text, page_count, word_count)
    """
    text = ""
    page_count = 0
    
    try:
        # Try pdfplumber first (better for complex layouts)
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                socketio.emit('extraction_progress', {
                    'current': i + 1,
                    'total': page_count,
                    'status': f'Extracting page {i + 1} of {page_count}'
                })
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                time.sleep(0.1)  # Small delay for progress visualization
    except Exception as e:
        # Fallback to PyPDF2
        try:
            reader = PdfReader(file_path)
            page_count = len(reader.pages)
            for i, page in enumerate(reader.pages):
                socketio.emit('extraction_progress', {
                    'current': i + 1,
                    'total': page_count,
                    'status': f'Extracting page {i + 1} of {page_count}'
                })
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                time.sleep(0.1)
        except Exception:
            return None, 0, 0
    
    word_count = len(text.split())
    return text.strip(), page_count, word_count



import google.generativeai as genai
from google.generativeai import GenerativeModel, ChatSession

genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

def summarize_text_gemini(text, summary_type='standard'):
    prompts = {
        'brief': "Provide a very concise 3-sentence summary of the following document.",
        'standard': "Summarize the following document into exactly 5 clear, informative bullet points.",
        'detailed': "Provide a comprehensive 8-point summary with detailed insights from this document.",
        'bullet': "Create a structured bullet-point summary with main topics and sub-points from this document.",
        'academic': "Provide an academic-style summary with introduction, key findings, methodology, and conclusion."
    }

    prompt = prompts.get(summary_type, prompts['standard'])
    full_prompt = f"{prompt}\n\n{text}"

    try:
        model_id = os.getenv("GEMINI_MODEL_ID", "gemini-2.5-flash")
        model = GenerativeModel(model_name=model_id)  # <-- FIXED HERE
        response = model.generate_content(full_prompt)
        return response.text
    except Exception as e:
        return f"Error during summarization: {str(e)}"




def create_pdf_summary(summary_text, filename):
    """Create a PDF file from summary text"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor='#2563eb',
        spaceAfter=30
    )
    story.append(Paragraph(f"Summary: {filename}", title_style))
    story.append(Spacer(1, 0.2 * inch))
    
    # Summary content
    for line in summary_text.split('\n'):
        if line.strip():
            story.append(Paragraph(line, styles['BodyText']))
            story.append(Spacer(1, 0.1 * inch))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def create_docx_summary(summary_text, filename):
    """Create a DOCX file from summary text"""
    doc = Document()
    
    # Title
    title = doc.add_heading(f'Summary: {filename}', 0)
    title.runs[0].font.color.rgb = RGBColor(37, 99, 235)
    
    # Summary content
    for line in summary_text.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line)
            p.runs[0].font.size = Pt(11)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==================== ROUTES ====================

@app.route('/')
def index():
    """Homepage with upload form"""
    return render_template('index.html')

@app.route('/summarize', methods=['POST'])
def summarize():
    """Handle PDF upload and summarization"""
    if 'pdf_file' not in request.files:
        return jsonify({'error': 'No PDF file uploaded'}), 400
    
    file = request.files['pdf_file']
    summary_type = request.form.get('summary_type', 'standard')
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'Invalid file type. Only PDF allowed'}), 400
    
    filename = secure_filename(file.filename)
    temp_path = os.path.join(UPLOAD_FOLDER, filename)
    
    try:
        file.save(temp_path)
        
        # Extract text from PDF
        socketio.emit('process_status', {'status': 'Extracting text from PDF...'})
        text, page_count, word_count = extract_text_from_pdf(temp_path)
        
        if not text or len(text.strip()) == 0:
            return jsonify({'error': 'Could not extract text from PDF'}), 400
        
        # Limit text length for API
        if len(text) > MAX_TEXT_LENGTH:
            text = text[:MAX_TEXT_LENGTH]
            socketio.emit('process_status', {
                'status': f'Text truncated to {MAX_TEXT_LENGTH} characters'
            })
        
        # Summarize text
        socketio.emit('process_status', {'status': 'Generating AI summary...'})
        summary = summarize_text_gemini(text, summary_type)
        
        # Save to history
        history_entry = SummaryHistory(
            filename=filename,
            summary_text=summary,
            summary_type=summary_type,
            page_count=page_count,
            word_count=word_count
        )
        db.session.add(history_entry)
        db.session.commit()
        
        socketio.emit('process_complete', {
            'status': 'Summary generated successfully!',
            'history_id': history_entry.id
        })
        
        return jsonify({
            'success': True,
            'summary': summary,
            'filename': filename,
            'page_count': page_count,
            'word_count': word_count,
            'summary_type': summary_type,
            'history_id': history_entry.id
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
    finally:
        # Clean up temporary file
        if os.path.exists(temp_path):
            os.remove(temp_path)

@app.route('/history')
def get_history():
    """Get summary history"""
    histories = SummaryHistory.query.order_by(SummaryHistory.created_at.desc()).limit(50).all()
    return jsonify([h.to_dict() for h in histories])

@app.route('/history/<int:history_id>')
def get_history_item(history_id):
    """Get specific history item"""
    history = SummaryHistory.query.get_or_404(history_id)
    return jsonify(history.to_dict())

@app.route('/history/<int:history_id>/rate', methods=['POST'])
def rate_summary(history_id):
    """Rate a summary"""
    history = SummaryHistory.query.get_or_404(history_id)
    rating = request.json.get('rating', 0)
    history.rating = max(1, min(5, int(rating)))  # Clamp between 1-5
    db.session.commit()
    return jsonify({'success': True, 'rating': history.rating})

@app.route('/history/<int:history_id>/delete', methods=['DELETE'])
def delete_history(history_id):
    """Delete history item"""
    history = SummaryHistory.query.get_or_404(history_id)
    db.session.delete(history)
    db.session.commit()
    return jsonify({'success': True})

@app.route('/download/<format>')
def download_summary(format):
    """Download summary in specified format"""
    summary_text = request.args.get('summary', '')
    filename = request.args.get('filename', 'summary')
    
    if format == 'txt':
        buffer = io.BytesIO()
        buffer.write(summary_text.encode('utf-8'))
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{filename}_summary.txt",
            mimetype='text/plain'
        )
    
    elif format == 'pdf':
        buffer = create_pdf_summary(summary_text, filename)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{filename}_summary.pdf",
            mimetype='application/pdf'
        )
    
    elif format == 'docx':
        buffer = create_docx_summary(summary_text, filename)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{filename}_summary.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    return jsonify({'error': 'Invalid format'}), 400

@app.route('/theme', methods=['POST'])
def set_theme():
    """Set user theme preference"""
    theme = request.json.get('theme', 'light')
    session['theme'] = theme
    return jsonify({'success': True, 'theme': theme})

@app.route('/stats')
def get_stats():
    """Get application statistics"""
    total_summaries = SummaryHistory.query.count()
    avg_rating = db.session.query(db.func.avg(SummaryHistory.rating)).scalar() or 0
    total_pages = db.session.query(db.func.sum(SummaryHistory.page_count)).scalar() or 0
    
    return jsonify({
        'total_summaries': total_summaries,
        'avg_rating': round(avg_rating, 2),
        'total_pages_processed': total_pages
    })

# ==================== WEBSOCKET EVENTS ====================

@socketio.on('connect')
def handle_connect():
    """Handle client connection"""
    emit('connected', {'data': 'Connected to server'})

@socketio.on('disconnect')
def handle_disconnect():
    """Handle client disconnection"""
    print('Client disconnected')

# ==================== RUN APPLICATION ====================

if __name__ == '__main__':
    socketio.run(app, debug=True, host='0.0.0.0', port=5001)

